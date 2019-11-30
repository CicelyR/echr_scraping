from datetime import datetime
import logging
import time
import unicodedata

from bs4 import BeautifulSoup
from docx import Document
from selenium import webdriver
from selenium.common.exceptions import (
    StaleElementReferenceException, TimeoutException)
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

DELAY = 5 # seconds
log = logging.Logger(__name__)

class Case:
    def __init__(self, application_no: str, browser: webdriver):
        self.application_no = application_no
        self.hudoc_url = ("https://hudoc.echr.coe.int/eng#{%22languageisocode%22:[%22ENG%22],"
                          f"%22appno%22:[%22{application_no}%22],"
                          "%22documentcollectionid2%22:[%22JUDGMENTS%22]}")
        self.exec_url = ("https://hudoc.exec.coe.int/eng#{%22EXECDocumentTypeCollection%22:[%22CEC%22],"
                f"%22EXECLanguage%22:[%22ENG%22],%22EXECAppno%22:[%22{application_no}%22]}}")
        self.cofm_url = ("https://hudoc.echr.coe.int/eng#{%22languageisocode%22:[%22ENG%22],"
                f"%22appno%22:[%22{application_no}%22],%22originatingbody%22:[%2217%22]}}")
        self.plan_url = ("https://hudoc.exec.coe.int/eng#{%22EXECDocumentTypeCollection%22:[%22acp%22,"
                f"%22acr%22],%22EXECAppno%22:[%22{application_no}%22],%22EXECLanguage%22:[%22ENG%22]}}")
        self.other_url = ("https://hudoc.exec.coe.int/eng#{%22EXECDocumentTypeCollection%22:"
                "[%22obs%22,%22CMDEC%22,%22CMINF%22,%22CMNOT%22,%22HEXEC%22],"
                f"%22EXECLanguage%22:[%22ENG%22],%22EXECAppno%22:[%22{application_no}%22]}}")
        self.judgements = self.get_documents(browser, self.hudoc_url)
        self.exec_cases = self.get_documents(browser, self.exec_url)
        self.cofm_documents = self.get_documents(browser, self.cofm_url)
        self.plan_documents = self.get_documents(browser, self.plan_url)
        self.other_documents = self.get_documents(browser, self.other_url)
        if self.validate_data():
            self.save_docx("case_docs/"+application_no.replace("/", "_")+".docx")
        else:
            log.warning(f"Validation failed for {application_no}")


    def get_documents(self, browser: webdriver, url: str):
        browser.get(url)

        try:
            WebDriverWait(browser, DELAY).until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div.headlineContaniner > a')))
        except TimeoutException:
            return []

        try:
            links = [link.get_attribute('href') for link in browser.find_elements_by_css_selector('div.headlineContaniner > a')]
        except StaleElementReferenceException:
            time.sleep(2)
            links = [link.get_attribute('href') for link in browser.find_elements_by_css_selector('div.headlineContaniner > a')]

        documents = []
        for link in links:
            documents.append(self.get_document_details(browser, link))
            
        return documents
        
    def get_document_details(self, browser: webdriver, link:str):
        case_details = {}
        browser.get(link)
        try:
            WebDriverWait(browser, DELAY).until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div#document > div.content p')))
        except TimeoutException:
            pass
        document_div = BeautifulSoup(browser.page_source, "lxml").select_one("div#document")
        case_details_tab = WebDriverWait(browser, DELAY).until(EC.presence_of_element_located((By.CSS_SELECTOR, 'li#notice > a')))
        case_details_tab.click()
        WebDriverWait(browser, DELAY).until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div#notice > div.content > div')))
        case_details_div = BeautifulSoup(browser.page_source, "lxml").select_one("div#notice")
        for div in case_details_div.select("div.row"):
            heading = div.select_one("div.noticefieldheading").text
            value = [d.text.strip() for d in div.select("div.noticefieldvalue div")]
            case_details[heading] = value
        document = "\n".join([unicodedata.normalize("NFKC", d.get_text()) for d in document_div.select("div.content div p")])
        return case_details, document
    
    def get_highest_court_judgement(self):
        if len(self.judgements) == 0:
            return None
        elif len(self.judgements) == 1:
            return self.judgements[0]
        else:
            judgement_dates = [datetime.strptime(j[0]["Judgment Date"][0], "%d/%M/%Y") for j in self.judgements]
            max_ix = judgement_dates.index(max(judgement_dates))
            return self.judgements[max_ix]

    def validate_data(self):
        if len(self.judgements) == 0:
            log.warning("Judgement missing for application %s", self.application_no)
            return False
        if len(self.judgements) > 1:
            log.warning("Multiple judgements for application %s", self.application_no)
            return False
        if len(self.exec_cases) > 1:
            log.warning("Multiple exec cases for application %s", self.application_no)
            return False
        return True

    def save_docx(self, filepath: str):
        doc = Document()

        doc.add_heading('HUDOC EXEC metadata', 1)

        if len(self.exec_cases) > 0:
            for key, value in self.exec_cases[0][0].items():
                doc.add_heading(key, 3)
                for item in value:
                    doc.add_paragraph(item)

        doc.add_heading('HUDOC EXEC case document', 1)
        if len(self.exec_cases) > 0:
            doc.add_paragraph(self.exec_cases[0][1])
        
        doc.add_heading('HUDOC metadata', 1)
        judgement = self.get_highest_court_judgement()
        if judgement is not None:
            for key, value in judgement[0].items():
                doc.add_heading(key, 3)
                for item in value:
                    doc.add_paragraph(item)
        else:
            doc.add_paragraph("No judgement found.")


        doc.add_heading("Committee of Minister resolutions/other documents", 1)
        for cofm_doc in self.cofm_documents:
            for key, value in cofm_doc[0].items():
                doc.add_heading(key, 3)
                for item in value:
                    doc.add_paragraph(item)
            doc.add_heading("Document", 3)
            doc.add_paragraph(cofm_doc[1])

        doc.add_heading("Action Plans")
        for plan_doc in self.plan_documents:
            for key, value in plan_doc[0].items():
                doc.add_heading(key, 3)
                for item in value:
                    doc.add_paragraph(item)
            doc.add_heading("Document", 3)
            doc.add_paragraph(plan_doc[1])

        doc.add_heading("HUDOC judgment (highest court)")
        if judgement is not None:
            doc.add_paragraph(judgement[1])
        else:
            doc.add_paragraph("No judgement found.")


        doc.add_heading("Other documents")
        for other_doc in self.other_documents:
            for key, value in other_doc[0].items():
                doc.add_heading(key, 3)
                for item in value:
                    doc.add_paragraph(item)
            doc.add_heading("Document", 3)
        
        doc.save(filepath)
